"""Orchestrateur du pipeline complet : audio -> transcription -> diarisation -> compte rendu.

Optimise pour les reunions longues (1-2h) :
- Gestion memoire : libere les ressources entre chaque etape
- Estimations de temps realistes
- Progression detaillee
"""

import gc
import json
import logging
import shutil
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Callable, List, Optional

from core.audio_processor import AudioProcessor
from core.transcriber import Transcriber, TranscriptionResult
from core.diarizer import DualChannelDiarizer
from core.summarizer import MeetingSummarizer
from config.settings import AppSettings

logger = logging.getLogger(__name__)


@dataclass
class PipelineResult:
    mic_wav: Path
    loopback_wav: Path
    mixed_wav: Path
    transcript_json: Path
    transcript_txt: Path
    minutes_md: Path
    minutes_docx: Optional[Path]
    transcript: TranscriptionResult
    minutes_text: str
    processing_time: float


def _estimate_processing_time(duration_sec: float, has_gpu: bool) -> str:
    """Estime le temps de traitement total."""
    # Whisper large-v3 : ~0.5x temps reel sur GPU, ~3-5x sur CPU
    if has_gpu:
        whisper_time = duration_sec * 0.5
    else:
        whisper_time = duration_sec * 4.0

    # LLM : ~2-5 min selon la longueur
    llm_time = 120 + (duration_sec / 3600) * 180

    # Audio processing : rapide
    audio_time = 30

    total = whisper_time + llm_time + audio_time
    if total < 120:
        return f"~{int(total / 60)} min"
    else:
        return f"~{int(total / 60)} min ({int(total / 3600)}h{int((total % 3600) / 60):02d})"


def _check_disk_space(output_dir: Path, required_mb: float = 500) -> bool:
    """Verifie qu'il y a assez d'espace disque."""
    try:
        usage = shutil.disk_usage(output_dir)
        free_mb = usage.free / (1024 * 1024)
        return free_mb >= required_mb
    except Exception:
        return True  # En cas d'erreur, continuer


class ProcessingPipeline:
    """Orchestre le traitement complet post-enregistrement."""

    def __init__(self, settings: AppSettings, transcriber: Optional[Transcriber] = None):
        self._settings = settings
        self._processor = AudioProcessor()
        self._transcriber = transcriber or Transcriber(
            model_size=settings.whisper_model,
            device=settings.whisper_device,
            compute_type=settings.whisper_compute_type,
            models_dir=settings.models_directory,
        )
        self._diarizer = DualChannelDiarizer(
            ratio_threshold=1.5,
            min_silence_duration=settings.min_silence_duration,
        )
        self._summarizer = MeetingSummarizer(
            model=settings.ollama_model,
            host=settings.ollama_host,
        )

    def process(self, mic_path: Path, loopback_path: Path, output_dir: Path,
                on_status: Optional[Callable[[str, float], None]] = None,
                on_token: Optional[Callable[[str], None]] = None,
                skip_llm: bool = False,
                existing_transcript: Optional[TranscriptionResult] = None,
                existing_chunk_summaries: Optional[List[str]] = None) -> PipelineResult:
        """Execute le pipeline complet.

        Args:
            existing_transcript: Si fourni, saute la transcription Whisper.
            existing_chunk_summaries: Si fourni, saute le resume par morceaux
                                      et fait directement la synthese finale.
        """
        start_time = time.time()

        def status(msg: str, progress: float):
            logger.info("[%.0f%%] %s", progress * 100, msg)
            if on_status:
                on_status(msg, progress)

        # Info sur l'enregistrement
        duration_sec = self._processor.get_wav_duration(mic_path)
        duration_min = duration_sec / 60

        # Verifier l'espace disque
        if not _check_disk_space(output_dir, required_mb=500):
            raise RuntimeError("Espace disque insuffisant (< 500 Mo libre)")

        # === ETAPE 1 : Preparation audio ===
        status(f"Preparation audio ({duration_min:.0f} min)...", 0.02)
        mixed_path = output_dir / "mix.wav"
        self._processor.mix_to_mono(mic_path, loopback_path, mixed_path)
        gc.collect()

        if existing_transcript and len(existing_transcript.segments) > 0:
            # === MODE RAPIDE : utiliser la transcription en direct ===
            logger.info("=== Pipeline rapide: %d segments deja transcrits ===",
                        len(existing_transcript.segments))
            status("Utilisation de la transcription en direct...", 0.60)
            transcript = existing_transcript
            # S'assurer que la duree est correcte
            if transcript.duration <= 0:
                transcript = TranscriptionResult(
                    language=transcript.language,
                    segments=transcript.segments,
                    duration=duration_sec,
                )
            n_segments = len(transcript.segments)
            logger.info("Transcription live: %d segments, %.0f min", n_segments, duration_min)
        else:
            # === MODE COMPLET : transcription Whisper ===
            has_gpu = self._settings.whisper_device in ("cuda", "auto")
            estimate = _estimate_processing_time(duration_sec, has_gpu)

            logger.info("=== Pipeline complet: %.0f min d'enregistrement ===", duration_min)
            logger.info("Temps estime: %s", estimate)
            status(f"Enregistrement de {duration_min:.0f} min — Temps estime: {estimate}", 0.0)

            # Chargement du modele Whisper
            if not self._transcriber.is_loaded:
                status("Chargement du modele Whisper (premiere utilisation)...", 0.08)
                self._transcriber.load_model()

            # Transcription (etape la plus longue)
            status(f"Transcription en cours ({duration_min:.0f} min d'audio)...", 0.12)

            def on_transcribe_progress(p: float):
                elapsed = time.time() - start_time
                if p > 0.05:
                    eta_total = elapsed / (0.12 + p * 0.48)
                    eta_remaining = eta_total - elapsed
                    eta_min = int(eta_remaining / 60)
                    status(f"Transcription... {p*100:.0f}% (~{eta_min} min restantes)", 0.12 + p * 0.48)
                else:
                    status(f"Transcription... {p*100:.0f}%", 0.12 + p * 0.48)

            transcript = self._transcriber.transcribe(
                mixed_path,
                language=self._settings.language,
                on_progress=on_transcribe_progress,
            )

            n_segments = len(transcript.segments)
            logger.info("Transcription: %d segments, %.0f min", n_segments, transcript.duration / 60)

        # === ETAPE 4 : Diarisation ===
        status(f"Identification des locuteurs ({n_segments} segments)...", 0.62)
        transcript = self._diarizer.diarize(mic_path, loopback_path, transcript)
        gc.collect()

        # === ETAPE 5 : Export transcription ===
        status("Export de la transcription...", 0.67)
        transcript_json_path = output_dir / "transcription.json"
        transcript_txt_path = output_dir / "transcription.txt"
        self._export_transcript_json(transcript, transcript_json_path)
        self._export_transcript_txt(transcript, transcript_txt_path)

        # === ETAPE 6 : Generation du compte rendu ===
        minutes_text = ""
        if not skip_llm:
            available, err_msg = self._summarizer.check_available()
            if available:
                if existing_chunk_summaries and len(existing_chunk_summaries) > 0:
                    # MODE RAPIDE : resumes partiels deja faits pendant l'enregistrement
                    logger.info("Synthese finale a partir de %d resumes existants",
                                len(existing_chunk_summaries))
                    status(f"Synthese finale ({len(existing_chunk_summaries)} resumes)...", 0.70)

                    from core.summarizer import (
                        SUMMARY_TEMPLATE, FINAL_SYNTHESIS_TEMPLATE, SYSTEM_PROMPT,
                    )
                    all_summaries = "\n\n".join(existing_chunk_summaries)
                    summary_section = SUMMARY_TEMPLATE.format(transcript_text=all_summaries)
                    final_prompt = FINAL_SYNTHESIS_TEMPLATE.format(
                        n_parts=len(existing_chunk_summaries),
                        duration_min=int(transcript.duration / 60),
                        summary_template=summary_section,
                    )
                    messages = [
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": final_prompt},
                    ]
                    client = self._summarizer._get_client()
                    minutes_text = self._summarizer._stream_response(
                        client, messages, on_token)
                else:
                    # MODE COMPLET : tout generer depuis la transcription
                    n_chunks_est = max(1, int(transcript.duration / 600))
                    status(f"Generation du compte rendu (~{n_chunks_est} morceaux)...", 0.70)

                    def on_chunk_progress(current: int, total: int):
                        chunk_progress = current / total
                        status(f"Resume morceau {current}/{total}...",
                               0.70 + chunk_progress * 0.15)

                    session_date = datetime.now().strftime("%d/%m/%Y")
                    minutes_text = self._summarizer.generate_minutes(
                        transcript,
                        session_date=session_date,
                        on_token=on_token,
                        on_chunk_progress=on_chunk_progress,
                    )
            else:
                logger.warning("Ollama non disponible: %s", err_msg)
                status(f"LLM indisponible: {err_msg}", 0.70)
                minutes_text = f"# Compte rendu non genere\n\nOllama non disponible : {err_msg}"
        else:
            minutes_text = "# Compte rendu\n\n_Generation LLM ignoree._"

        # === ETAPE 7 : Export compte rendu ===
        status("Export du compte rendu...", 0.92)
        minutes_md_path = output_dir / "compte_rendu.md"
        with open(minutes_md_path, "w", encoding="utf-8") as f:
            f.write(minutes_text)

        minutes_docx_path = None
        try:
            minutes_docx_path = output_dir / "compte_rendu.docx"
            self._export_minutes_docx(minutes_text, transcript, minutes_docx_path)
        except Exception as e:
            logger.warning("Export DOCX echoue: %s", e)
            minutes_docx_path = None

        processing_time = time.time() - start_time
        proc_min = int(processing_time / 60)
        proc_sec = int(processing_time % 60)
        status(f"Termine ! (traitement: {proc_min}min {proc_sec}s)", 1.0)
        logger.info("=== Pipeline termine en %dmin %ds pour %.0f min d'audio ===",
                     proc_min, proc_sec, duration_min)

        return PipelineResult(
            mic_wav=mic_path,
            loopback_wav=loopback_path,
            mixed_wav=mixed_path,
            transcript_json=transcript_json_path,
            transcript_txt=transcript_txt_path,
            minutes_md=minutes_md_path,
            minutes_docx=minutes_docx_path,
            transcript=transcript,
            minutes_text=minutes_text,
            processing_time=processing_time,
        )

    def _export_transcript_json(self, transcript: TranscriptionResult, path: Path):
        """Sauvegarde la transcription en JSON structure."""
        data = {
            "language": transcript.language,
            "duration_seconds": transcript.duration,
            "segments": [
                {
                    "start": seg.start,
                    "end": seg.end,
                    "speaker": seg.speaker,
                    "text": seg.text,
                    "words": [
                        {
                            "start": w.start,
                            "end": w.end,
                            "word": w.word,
                            "probability": round(w.probability, 3),
                        }
                        for w in seg.words
                    ],
                }
                for seg in transcript.segments
            ],
        }
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    def _export_transcript_txt(self, transcript: TranscriptionResult, path: Path):
        """Sauvegarde la transcription en texte lisible."""
        duration_min = int(transcript.duration) // 60
        duration_sec = int(transcript.duration) % 60

        lines = [
            f"Transcription de la reunion",
            f"Duree : {duration_min}min {duration_sec}sec",
            "=" * 50,
            "",
        ]

        for seg in transcript.segments:
            speaker = seg.speaker or "Inconnu"
            minutes = int(seg.start) // 60
            seconds = int(seg.start) % 60
            lines.append(f"[{minutes:02d}:{seconds:02d}] {speaker} :")
            lines.append(f"{seg.text}")
            lines.append("")

        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _export_minutes_docx(self, minutes_md: str, transcript: TranscriptionResult, path: Path):
        """Convertit le compte rendu MD en DOCX formate."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = doc.add_heading("Compte Rendu de Reunion", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_str = datetime.now().strftime("%d/%m/%Y %H:%M")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Date : {date_str}")
        run.font.size = Pt(10)
        run.italic = True

        if transcript:
            duration_min = int(transcript.duration / 60)
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(f"Duree : {duration_min} minutes")
            run2.font.size = Pt(10)
            run2.italic = True

        doc.add_paragraph()

        for line in minutes_md.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                continue
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip("*"), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip("*"), level=2)
            elif line.startswith("| ") and "---" not in line:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                doc.add_paragraph("  |  ".join(cells))
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line[0].isdigit() and ". " in line[:5]:
                idx = line.index(". ")
                doc.add_paragraph(line[idx + 2:], style="List Number")
            else:
                clean = line.replace("**", "").replace("__", "")
                doc.add_paragraph(clean)

        doc.save(str(path))
        logger.info("DOCX exporte: %s", path.name)
